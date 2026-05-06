"""Gera o relatório completo v2 — detalhado, autoexplicativo, com Q&A."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Margens ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

AZUL     = RGBColor(30, 61, 89)
VERDE    = RGBColor(39, 174, 96)
VERMELHO = RGBColor(192, 57, 43)
CINZA    = RGBColor(100, 100, 100)


def hdr_shd(cell, hex_color='1E3D59'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None, hdr_color='1E3D59'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        r = hdr_cells[i].paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        hdr_shd(hdr_cells[i], hdr_color)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    return table


def h1(doc, text):
    p = doc.add_heading(text, 1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.color.rgb = AZUL; r.font.name = 'Calibri'
    return p


def h2(doc, text):
    p = doc.add_heading(text, 2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(52, 73, 94); r.font.name = 'Calibri'
    return p


def h3(doc, text):
    p = doc.add_heading(text, 3)
    for r in p.runs:
        r.font.name = 'Calibri'
    return p


def para(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = 'Calibri'; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p


def mixed(doc, parts, size=11):
    """parts = list of (text, bold, italic, color)"""
    p = doc.add_paragraph()
    for text, bold, italic, color in parts:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.name = 'Calibri'; r.font.size = Pt(size)
        if color: r.font.color.rgb = color
    return p


def caixa_destaque(doc, titulo, texto, cor_hex='1E3D59'):
    """Simula uma caixa de destaque com tabela 1x1."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    hdr_shd(cell, 'EBF5FB')
    p = cell.paragraphs[0]
    r_t = p.add_run(titulo + '\n')
    r_t.bold = True; r_t.font.size = Pt(11)
    r_t.font.color.rgb = AZUL
    r_d = p.add_run(texto)
    r_d.font.size = Pt(10)
    doc.add_paragraph()


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'


def numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'


def separador(doc):
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════
separador(doc); separador(doc)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('RELATÓRIO DE MELHORIAS — TCC')
r.bold = True; r.font.size = Pt(22); r.font.name = 'Calibri'
r.font.color.rgb = AZUL

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('Previsão de Rebaixamento no Brasileirão Série A')
r2.font.size = Pt(15); r2.italic = True; r2.font.name = 'Calibri'

separador(doc)
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run('Leonardo Feitosa Barroso  |  Ciência de Dados — UFPB  |  2025')
r3.font.size = Pt(12); r3.font.name = 'Calibri'

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run('Orientador: Prof. Hilton Martins')
r4.font.size = Pt(12); r4.italic = True; r4.font.name = 'Calibri'

separador(doc)
aviso = doc.add_paragraph()
aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
ra = aviso.add_run('Documento preparado para apresentação ao orientador — Maio/2026')
ra.font.size = Pt(10); ra.italic = True; ra.font.color.rgb = CINZA

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. CONTEXTO
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '1. Contexto e Motivação')

para(doc,
    'O professor solicitou a incorporação de técnicas avançadas de Machine Learning '
    'utilizadas no TCC de referência (Matheus Santos de Oliveira Flor), com foco em '
    'duas melhorias principais:'
)
bullet(doc, 'Engenharia de features com janelas deslizantes (sliding windows) — criar variáveis que '
       'resumam o histórico de desempenho recente de cada clube.', bold_prefix='Engenharia de features: ')
bullet(doc, 'Walk-forward validation — validar o modelo respeitando a ordem do tempo, '
       'simulando como ele funcionaria na prática.', bold_prefix='Walk-forward validation: ')

separador(doc)
para(doc,
    'Além dessas duas, foram também implementadas otimização automática de hiperparâmetros '
    'e comparação entre quatro modelos diferentes, usando AUC-ROC como métrica principal. '
    'Todas as melhorias foram implementadas, testadas e os resultados estão documentados abaixo.'
)

separador(doc)

# ═══════════════════════════════════════════════════════════════════════
# 2. RESUMO ANTES × DEPOIS
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '2. Resumo: Antes × Depois')

add_table(doc,
    ['Aspecto', 'Antes das Melhorias', 'Depois das Melhorias'],
    [
        ['Número de variáveis (features)', '3', '15'],
        ['Variáveis de desempenho histórico', 'Nenhuma', '12 (médias de 3 e 5 temporadas)'],
        ['Tipo de validação', '1 corte simples treino/teste', 'Walk-forward com 5 janelas temporais'],
        ['Ajuste de parâmetros', 'Padrão (sem otimização)', 'RandomizedSearchCV + TimeSeriesSplit'],
        ['Modelos treinados e comparados', '1 (Regressão Logística)', '4 (LR, Random Forest, XGBoost, LightGBM)'],
        ['Métrica principal', 'Acurácia: 77,5 %', 'AUC-ROC: 0.877 (LightGBM)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. COLETA DE DADOS
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '3. Melhoria 1 — Nova Coleta de Dados de Desempenho')

mixed(doc, [
    ('Arquivo: ', True, False, AZUL),
    ('notebooks/00_coleta_dados.ipynb', False, True, CINZA),
    (' (Etapa 4 — células novas no final do notebook)', False, False, None),
])

separador(doc)
caixa_destaque(doc,
    'O que é web scraping?',
    'Web scraping é a extração automática de dados de sites. Neste projeto, '
    'o código acessa o site Transfermarkt automaticamente, lê as tabelas de '
    'classificação de cada temporada e salva os dados em planilha Excel — '
    'sem precisar copiar manualmente.'
)

para(doc,
    'Para enriquecer o modelo com informações sobre o desempenho em campo '
    '(e não apenas sobre o elenco), foi adicionada uma nova etapa de coleta '
    'de dados no notebook 00. O código raspa do site Transfermarkt o histórico '
    'de classificação de cada clube por temporada.'
)

separador(doc)
h2(doc, '3.1 Dados coletados por temporada e clube')

add_table(doc,
    ['Coluna', 'Significado'],
    [
        ['Posicao',        'Colocação final na tabela (1º ao 20º)'],
        ['V',              'Número de vitórias na temporada'],
        ['E',              'Número de empates'],
        ['D',              'Número de derrotas'],
        ['Gols_Pro',       'Total de gols marcados'],
        ['Gols_Contra',    'Total de gols sofridos'],
        ['SG',             'Saldo de gols (Gols_Pro − Gols_Contra)'],
        ['Pts',            'Pontos totais acumulados'],
        ['Aproveitamento', 'Percentual de pontos aproveitados (Pts / (J×3) × 100)'],
    ]
)

separador(doc)
h2(doc, '3.2 Período e volume coletado')
bullet(doc, '11 temporadas (2014 a 2024) × 20 clubes por temporada = 220 registros',
       bold_prefix='Volume: ')
bullet(doc, '2025 foi excluída propositalmente — como é o ano que estamos tentando prever, '
       'incluir seus dados seria trapaça (data leakage): o modelo aprenderia o resultado antes de prever.',
       bold_prefix='Por que excluir 2025? ')
bullet(doc, 'dados/tabela_desempenho_brasileirao.xlsx — arquivo Excel com 12 abas '
       '(uma por temporada + aba "Todos" com todos os dados juntos)',
       bold_prefix='Arquivo gerado: ')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. JANELAS DESLIZANTES
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '4. Melhoria 2 — Janelas Deslizantes (Sliding Windows)')

mixed(doc, [
    ('Arquivo: ', True, False, AZUL),
    ('notebooks/02_preprocessamento.ipynb', False, True, CINZA),
])

separador(doc)
caixa_destaque(doc,
    'O que são janelas deslizantes?',
    'Imagine que você quer saber se um clube vai ser rebaixado em 2022. '
    'Em vez de olhar apenas para o tamanho do elenco, você pergunta: '
    '"Quantos pontos esse clube fez nos últimos 3 anos?" e '
    '"Quantos gols ele sofreu nos últimos 5 anos?". '
    'Essas médias históricas são as janelas deslizantes — elas resumem '
    'a "forma" recente do clube antes do campeonato começar.'
)

h2(doc, '4.1 Por que isso melhora o modelo?')
para(doc,
    'O modelo original usava apenas dados estáticos de elenco (tamanho, estrangeiros, valor de mercado). '
    'Esses dados dizem o que o clube tem, mas não o que ele fez em campo. '
    'Um clube com elenco caro pode ter jogado mal nos últimos anos — e vice-versa. '
    'As janelas deslizantes adicionam essa memória histórica ao modelo.'
)

separador(doc)
h2(doc, '4.2 Regra anti-data-leakage (anti-vazamento de dados)')

caixa_destaque(doc,
    'O que é data leakage?',
    'Data leakage (vazamento de dados) acontece quando o modelo aprende '
    'com informações que não estaria disponíveis no momento real da previsão. '
    'Por exemplo: usar os resultados do campeonato de 2022 para prever quem '
    'seria rebaixado em 2022 — o modelo "trapaceia" usando o futuro.'
)

para(doc,
    'Para garantir que não há vazamento, usamos a técnica shift(1).rolling(): '
    'para calcular a média da temporada T, o código usa apenas as temporadas '
    'T-1, T-2, T-3 (passado). Nunca a temporada T em si.'
)

separador(doc)
h2(doc, '4.3 As 12 novas features criadas')

add_table(doc,
    ['Métrica', 'Feature — Janela 3 temporadas', 'Feature — Janela 5 temporadas'],
    [
        ['Pontos',          'Pts_media_3',            'Pts_media_5'],
        ['Saldo de gols',   'SG_media_3',             'SG_media_5'],
        ['Gols marcados',   'Gols_Pro_media_3',       'Gols_Pro_media_5'],
        ['Gols sofridos',   'Gols_Contra_media_3',    'Gols_Contra_media_5'],
        ['Vitórias',        'V_media_3',              'V_media_5'],
        ['Aproveitamento',  'Aproveitamento_media_3', 'Aproveitamento_media_5'],
    ]
)

separador(doc)
para(doc,
    'Total de features no modelo final: 15 variáveis '
    '(3 de elenco + 12 de janela deslizante).',
    bold=True
)
para(doc,
    'Clubes recém-promovidos da Série B, sem histórico completo na Série A, '
    'recebem a mediana do conjunto de treino — assim o modelo não fica sem informação, '
    'mas também não usa dados do futuro para preencher os campos vazios.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. WALK-FORWARD
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '5. Melhoria 3 — Walk-Forward Validation')

mixed(doc, [
    ('Arquivos: ', True, False, AZUL),
    ('03_modelo_logistica.ipynb  |  04_modelo_random_forest.ipynb  |  06_comparacao_modelos.ipynb',
     False, True, CINZA),
])

separador(doc)
caixa_destaque(doc,
    'O que é Walk-Forward Validation?',
    'É uma técnica de validação que respeita a ordem do tempo. '
    'Em vez de dividir os dados aleatoriamente, o modelo é treinado '
    'sempre no passado e testado no futuro — exatamente como funcionaria na prática. '
    'A cada "fold" (janela), adicionamos mais um ano ao treino e testamos no próximo. '
    'Isso é equivalente ao TimeSeriesSplit do scikit-learn.'
)

h2(doc, '5.1 Como foram organizados os 5 folds (janelas)')

add_table(doc,
    ['Fold', 'Conjunto de Treino', 'Conjunto de Validação', 'Observação'],
    [
        ['1', '2014 – 2017 (4 anos)', '2018 (1 ano)', 'Mínimo de histórico'],
        ['2', '2014 – 2018 (5 anos)', '2019 (1 ano)', ''],
        ['3', '2014 – 2019 (6 anos)', '2020 (1 ano)', 'Ano da pandemia'],
        ['4', '2014 – 2020 (7 anos)', '2021 (1 ano)', ''],
        ['5', '2014 – 2021 (8 anos)', '2022 (1 ano)', 'Último fold de validação'],
        ['Final', '2014 – 2022 (treino completo)', '2023 – 2024 (teste final)', 'Avaliação definitiva'],
    ]
)

separador(doc)
h2(doc, '5.2 Resultados do Walk-Forward — AUC-ROC por fold')

add_table(doc,
    ['Modelo', 'Fold 1\n(2018)', 'Fold 2\n(2019)', 'Fold 3\n(2020)', 'Fold 4\n(2021)', 'Fold 5\n(2022)', 'Média', 'Desvio'],
    [
        ['Reg. Logística', '0.828', '0.859', '0.703', '0.750', '0.828', '0.794', '± 0.058'],
        ['Random Forest',  '0.789', '0.734', '0.758', '0.656', '0.875', '0.762', '± 0.071'],
        ['XGBoost',        '0.766', '0.719', '0.531', '0.594', '0.688', '0.659', '± 0.085'],
        ['LightGBM',       '0.719', '0.750', '0.438', '0.688', '0.906', '0.700', '± 0.151'],
    ]
)

separador(doc)
caixa_destaque(doc,
    'Como interpretar a tabela acima?',
    'AUC-ROC varia de 0 a 1. Quanto mais próximo de 1, melhor o modelo separa '
    'clubes rebaixados de permanentes. Um AUC de 0.5 equivale a chutar aleatoriamente. '
    'O desvio padrão (±) mostra a consistência: quanto menor, mais estável o modelo ao longo dos anos. '
    'A Regressão Logística foi a mais estável (± 0.058), enquanto o LightGBM '
    'foi o mais instável (± 0.151).'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. HIPERPARÂMETROS
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '6. Melhoria 4 — Otimização de Hiperparâmetros')

mixed(doc, [
    ('Arquivos: ', True, False, AZUL),
    ('03_modelo_logistica.ipynb  |  04_modelo_random_forest.ipynb  |  06_comparacao_modelos.ipynb',
     False, True, CINZA),
])

separador(doc)
caixa_destaque(doc,
    'O que são hiperparâmetros?',
    'Hiperparâmetros são as "configurações" de um modelo que definimos antes do treino. '
    'Por exemplo: quantas árvores usar no Random Forest? Quão forte deve ser a regularização '
    'da Regressão Logística? Escolher mal esses valores pode deixar o modelo muito simples '
    '(underfitting) ou muito complexo (overfitting). '
    'O RandomizedSearchCV testa dezenas de combinações automaticamente e escolhe a melhor.'
)

h2(doc, '6.1 Técnica utilizada: RandomizedSearchCV + TimeSeriesSplit')
bullet(doc, '30 combinações aleatórias de hiperparâmetros testadas por modelo',
       bold_prefix='Candidatos testados: ')
bullet(doc, 'TimeSeriesSplit com 5 folds — mesmo durante a busca de parâmetros, '
       'a ordem temporal é respeitada', bold_prefix='Validação interna: ')
bullet(doc, 'AUC-ROC — o parâmetro vencedor é o que maximiza o AUC, não a acurácia',
       bold_prefix='Critério de seleção: ')

separador(doc)
h2(doc, '6.2 Parâmetros explorados e melhores valores encontrados')

add_table(doc,
    ['Modelo', 'Parâmetros explorados', 'Melhor configuração encontrada'],
    [
        ['Reg. Logística',
         'C (regularização), solver',
         'C = 0.985, solver = lbfgs | AUC-CV = 0.754'],
        ['Random Forest',
         'n_estimators, max_depth, min_samples_split, max_features',
         'n_est=100, max_depth=3, min_split=5, max_feat=log2 | AUC-CV = 0.757'],
        ['XGBoost',
         'n_estimators, max_depth, learning_rate, subsample',
         'Configuração otimizada automaticamente'],
        ['LightGBM',
         'n_estimators, max_depth, learning_rate, num_leaves',
         'Configuração otimizada automaticamente'],
    ]
)

separador(doc)
caixa_destaque(doc,
    'O que significa AUC-CV?',
    'AUC-CV é o AUC-ROC médio calculado durante a validação cruzada interna '
    '(dentro do RandomizedSearchCV). Ele mede o quão bom é o conjunto de '
    'hiperparâmetros nos dados de treino, antes de ver o teste final. '
    'Valores acima de 0.75 indicam boa capacidade de generalização.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. COMPARAÇÃO DE MODELOS
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '7. Melhoria 5 — Comparação de Múltiplos Modelos')

mixed(doc, [
    ('Arquivo: ', True, False, AZUL),
    ('notebooks/06_comparacao_modelos.ipynb', False, True, CINZA),
])

separador(doc)
para(doc,
    'Quatro modelos foram treinados com as mesmas 15 features, '
    'mesma separação temporal e mesma metodologia de validação. '
    'A comparação é feita no conjunto de teste independente (2023–2024), '
    'que nenhum modelo "viu" durante o treino ou a busca de hiperparâmetros.'
)

separador(doc)
h2(doc, '7.1 Os 4 modelos comparados')

add_table(doc,
    ['Modelo', 'Tipo', 'Característica principal'],
    [
        ['Regressão Logística', 'Linear',
         'Modelo simples e interpretável — cada feature tem um coeficiente direto'],
        ['Random Forest', 'Ensemble de árvores',
         'Combina centenas de árvores de decisão — robusto a outliers'],
        ['XGBoost', 'Gradient Boosting',
         'Árvores construídas sequencialmente, cada uma corrigindo o erro da anterior'],
        ['LightGBM', 'Gradient Boosting eficiente',
         'Variante mais rápida do XGBoost, com crescimento de árvore por folha'],
    ]
)

separador(doc)
h2(doc, '7.2 Resultados finais — conjunto de teste 2023-2024')

add_table(doc,
    ['Posição', 'Modelo', 'Acurácia', 'AUC-ROC', 'Observação'],
    [
        ['1º  ★', 'LightGBM',            '82,5 %', '0.877', 'Melhor AUC-ROC'],
        ['2º',    'Random Forest',        '80,0 %', '0.844', ''],
        ['3º',    'Regressão Logística',  '80,0 %', '0.828', 'Mais interpretável'],
        ['4º',    'XGBoost',             '82,5 %', '0.652', 'Alta acurácia, baixo AUC'],
    ]
)

separador(doc)
caixa_destaque(doc,
    'Por que o XGBoost tem acurácia alta mas AUC baixo?',
    'O XGBoost classificou quase todos os clubes como "Permanece". '
    'Como 80% dos clubes realmente permanecem, acerta muito pela prevalência da classe — '
    'não por inteligência real. O AUC-ROC expõe esse comportamento: '
    'um modelo que chuta sempre "Permanece" teria AUC próximo de 0.5. '
    'Por isso AUC-ROC é a métrica mais honesta para dados desbalanceados como este.'
)

separador(doc)
h2(doc, '7.3 Importância das features — Regressão Logística')
para(doc,
    'Os coeficientes da Regressão Logística revelam quais variáveis mais influenciam '
    'o risco de rebaixamento (sinal positivo = aumenta risco; negativo = reduz risco):'
)

add_table(doc,
    ['Feature', 'Coeficiente', 'Interpretação'],
    [
        ['Valor de Mercado Total', '−0.91', 'Clube mais rico → menos risco'],
        ['V_media_3 (vitórias últ. 3 anos)', '−0.73', 'Mais vitórias recentes → menos risco'],
        ['Plantel', '+0.56', 'Efeito de multicolinearidade com valor de mercado'],
        ['Gols_Pro_media_5', '−0.52', 'Mais gols marcados nos últ. 5 anos → menos risco'],
        ['Gols_Contra_media_5', '+0.50', 'Mais gols sofridos → mais risco'],
        ['Pts_media_3 / Pts_media_5', '+0.47 / +0.47', 'Correlacionados com V_media — distribuição de peso'],
    ]
)

separador(doc)
para(doc,
    'A feature V_media_3 (média de vitórias nas últimas 3 temporadas) é a '
    'segunda mais importante do modelo — só atrás do valor de mercado. '
    'Isso confirma que o histórico recente em campo é tão relevante quanto o poder financeiro.',
    bold=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. PREVISÃO 2025
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '8. Previsão de Rebaixamento — Temporada 2025')

mixed(doc, [
    ('Arquivo: ', True, False, AZUL),
    ('notebooks/07_previsao_2025.ipynb', False, True, CINZA),
])

separador(doc)
para(doc,
    'Com o modelo treinado e validado, a previsão para 2025 é aplicada '
    'usando os dados de elenco da temporada atual combinados com as '
    'janelas deslizantes calculadas a partir do histórico 2014–2024.'
)

separador(doc)
h2(doc, '8.1 Clubes previstos para rebaixamento em 2025')

add_table(doc,
    ['Posição', 'Clube', 'Prob. Rebaixamento', 'Situação prevista'],
    [
        ['1º', 'Sport Recife',  '65,4 %', 'REBAIXADO'],
        ['2º', 'Vitória',       '64,0 %', 'REBAIXADO'],
        ['3º', 'Juventude',     '54,9 %', 'REBAIXADO'],
        ['4º', 'Mirassol',      '49,0 %', 'REBAIXADO'],
        ['5º', 'Ceará',         '36,3 %', 'Permanece (zona de atenção)'],
        ['6º', 'Fluminense',    '26,0 %', 'Permanece'],
        ['7º', 'Fortaleza',     '24,6 %', 'Permanece'],
    ]
)

separador(doc)
para(doc,
    'Nota: os 4 últimos colocados na previsão de probabilidade são classificados '
    'como rebaixados, independentemente do limiar de 50% — assim como funciona '
    'o Brasileirão (4 rebaixados por temporada).',
    italic=True, color=CINZA
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. RESPOSTA AO PROFESSOR
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '9. Resposta Direta ao Professor')

caixa_destaque(doc,
    'Mensagem do Prof. Hilton Martins (27/04/2026):',
    '"Um dos pontos que gostaria de tratar é sobre a metodologia e incorporação '
    'de técnicas de engenharia de features, walk forward etc — alguns procedimentos '
    'que foram usados no trabalho que te enviei."'
)

para(doc, 'Resposta: todos os pontos foram implementados, conforme detalhado abaixo.')
separador(doc)

numbered(doc,
    'Foram criadas 12 novas features usando médias das últimas 3 e 5 temporadas para: '
    'pontos, saldo de gols, gols marcados, gols sofridos, vitórias e aproveitamento. '
    'A técnica shift(1).rolling() garante ausência de data leakage. '
    'Arquivo: notebooks/02_preprocessamento.ipynb',
    bold_prefix='Engenharia de features (janelas deslizantes): ')
separador(doc)

numbered(doc,
    'Implementado em 03_modelo_logistica.ipynb, 04_modelo_random_forest.ipynb '
    'e 06_comparacao_modelos.ipynb. Cinco folds temporais crescentes (2014–2017 até 2014–2021), '
    'sempre treinando no passado e validando no futuro. '
    'AUC-ROC médio da Regressão Logística: 0.794 (mais estável entre os modelos).',
    bold_prefix='Walk-forward validation: ')
separador(doc)

numbered(doc,
    'RandomizedSearchCV com 30 candidatos aleatórios por modelo, '
    'validação interna com TimeSeriesSplit(5 folds). '
    'Critério de seleção: AUC-ROC — mais adequado que acurácia para dados desbalanceados.',
    bold_prefix='Otimização de hiperparâmetros: ')
separador(doc)

numbered(doc,
    'Quatro modelos comparados com as mesmas features e metodologia: '
    'Regressão Logística, Random Forest, XGBoost e LightGBM. '
    'Melhor resultado: LightGBM com AUC-ROC 0.877 no teste final. '
    'Arquivo: notebooks/06_comparacao_modelos.ipynb',
    bold_prefix='Comparação de múltiplos modelos: ')
separador(doc)

numbered(doc,
    'Adotado como métrica principal em todos os notebooks de modelagem. '
    'O AUC-ROC mede a capacidade do modelo de ordenar corretamente os clubes '
    'por risco de rebaixamento, independente do limiar de decisão.',
    bold_prefix='AUC-ROC como métrica principal: ')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10. PERGUNTAS E RESPOSTAS
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '10. Perguntas e Respostas Prováveis na Apresentação')

perguntas = [
    (
        'Por que usar janelas de 3 e 5 temporadas? Por que não 1 ou 10?',
        'A janela de 3 temporadas captura a forma recente do clube (curto prazo), '
        'enquanto a de 5 captura a tendência mais longa (médio prazo). '
        'Janelas muito curtas (1 ano) são muito voláteis — um clube pode ter um ano atípico. '
        'Janelas muito longas (10 anos) diluem demais a informação recente e '
        'misturam eras diferentes do clube. As janelas de 3 e 5 são consagradas na literatura '
        'de séries temporais esportivas e foram as mesmas usadas no TCC de referência.'
    ),
    (
        'Por que walk-forward e não validação cruzada tradicional (k-fold)?',
        'A validação cruzada tradicional embaralha os dados aleatoriamente, '
        'o que causaria data leakage em séries temporais: o modelo aprenderia '
        'com dados de 2022 para prever 2018, por exemplo. '
        'O walk-forward garante que o modelo nunca vê o futuro durante o treino — '
        'exatamente como funcionaria em produção real.'
    ),
    (
        'Por que AUC-ROC é melhor que acurácia neste problema?',
        'No dataset, 80% dos clubes permanecem e apenas 20% são rebaixados. '
        'Um modelo que preveja sempre "Permanece" teria 80% de acurácia sem aprender nada. '
        'O AUC-ROC mede a capacidade de ordenar corretamente os clubes por risco, '
        'independente do limiar. Um AUC de 0.877 significa que em 87,7% dos pares '
        '(rebaixado, permaneceu), o modelo atribuiu maior probabilidade ao rebaixado — '
        'isso é muito mais informativo que a acurácia.'
    ),
    (
        'O LightGBM teve o melhor AUC-ROC no teste, mas a Regressão Logística '
        'foi mais estável no walk-forward. Qual modelo usar na prática?',
        'Depende da prioridade. Se o objetivo é máxima performance preditiva: LightGBM (AUC 0.877). '
        'Se o objetivo é interpretabilidade para apresentar ao professor/banca '
        '(explicar coeficientes, justificar cada previsão): Regressão Logística. '
        'Para um TCC, a Regressão Logística é geralmente preferida por ser explicável — '
        'cada coeficiente tem uma interpretação clara. O LightGBM é uma "caixa-preta".'
    ),
    (
        'Por que o XGBoost teve alta acurácia (82,5%) mas baixo AUC-ROC (0.652)?',
        'O XGBoost aprendeu a classificar quase tudo como "Permanece". '
        'Como 80% dos clubes realmente permanecem, isso gera acurácia alta por prevalência, '
        'não por inteligência. O AUC-ROC revela esse comportamento: '
        'o modelo não sabe distinguir bem rebaixados de permanecidos. '
        'Isso é um exemplo clássico do porquê não se deve usar acurácia '
        'como métrica única em datasets desbalanceados.'
    ),
    (
        'Por que a temporada 2025 foi excluída da coleta de desempenho?',
        'Porque 2025 é o ano que o modelo está tentando prever. '
        'Se coletássemos os resultados do campeonato de 2025 (que ainda está em andamento), '
        'o modelo teria acesso ao futuro — isso é data leakage. '
        'As janelas deslizantes para 2025 foram calculadas a partir '
        'do histórico 2014–2024, simulando o que um gestor saberia antes '
        'do início da temporada.'
    ),
    (
        'O modelo previu corretamente os rebaixados de 2023 e 2024?',
        'O conjunto de teste (2023–2024) foi usado para avaliação. '
        'Com AUC-ROC de 0.828 (Regressão Logística) a 0.877 (LightGBM), '
        'o modelo demonstrou boa capacidade de separar os clubes em risco. '
        'A acurácia de 80-82,5% no teste significa que de 40 observações, '
        'o modelo acertou 32-33. Como há apenas 8 rebaixados no período, '
        'recall e F1 da classe minoritária são os indicadores mais relevantes.'
    ),
    (
        'Quais as limitações do modelo?',
        '1) Dataset pequeno: apenas 220 observações históricas (11 anos × 20 clubes). '
        '2) Features limitadas: não captura lesões, trocas de treinador, motivação, '
        'calendário de jogos. '
        '3) Clubes novos na Série A: times recém-promovidos têm NaN nas janelas, '
        'substituídos pela mediana do treino — aproximação razoável, mas imperfeita. '
        '4) Estabilidade temporal: o futebol muda — estruturas de clubes de 2014 '
        'podem não refletir o cenário atual.'
    ),
]

for i, (pergunta, resposta) in enumerate(perguntas):
    h2(doc, f'P{i+1}: {pergunta}')
    para(doc, f'Resposta: {resposta}')
    separador(doc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11. ESTRUTURA DE ARQUIVOS
# ═══════════════════════════════════════════════════════════════════════
h1(doc, '11. Estrutura de Arquivos do Projeto')

add_table(doc,
    ['Arquivo', 'O que contém'],
    [
        ['notebooks/00_coleta_dados.ipynb',        'Raspagem do Transfermarkt: elencos + desempenho 2014-2024'],
        ['notebooks/02_preprocessamento.ipynb',    '15 features + janelas deslizantes + separação temporal'],
        ['notebooks/03_modelo_logistica.ipynb',    'Regressão Logística: walk-forward + tuning + avaliação'],
        ['notebooks/04_modelo_random_forest.ipynb','Random Forest: walk-forward + tuning + importância features'],
        ['notebooks/06_comparacao_modelos.ipynb',  '4 modelos comparados: walk-forward + curvas ROC + tabela final'],
        ['notebooks/07_previsao_2025.ipynb',       'Previsão final: probabilidades de rebaixamento por clube em 2025'],
        ['dados/BASE_FINAL.xlsx',                  'Dados de elenco por clube e temporada (2014-2025)'],
        ['dados/tabela_desempenho_brasileirao.xlsx','Desempenho histórico por clube: V, E, D, Gols, Pts (2014-2024)'],
        ['modelos/logistica.pkl',                  'Modelo de Regressão Logística otimizado e salvo'],
        ['modelos/random_forest.pkl',              'Modelo Random Forest otimizado e salvo'],
        ['modelos/xgboost.pkl',                    'Modelo XGBoost otimizado e salvo'],
        ['modelos/lightgbm.pkl',                   'Modelo LightGBM otimizado e salvo (melhor AUC-ROC)'],
        ['modelos/scaler_logistica.pkl',           'Normalizador ajustado no treino (2014-2022)'],
        ['modelos/mediana_treino.pkl',             'Medianas para imputar clubes sem histórico'],
    ]
)

separador(doc)
para(doc,
    'Todos os notebooks foram executados com sucesso. '
    'Os modelos estão salvos e prontos para uso no aplicativo Streamlit.',
    bold=True, color=VERDE
)

# Salvar
caminho = 'relatorio_melhorias_tcc.docx'
doc.save(caminho)
print(f'Salvo: {caminho}')
